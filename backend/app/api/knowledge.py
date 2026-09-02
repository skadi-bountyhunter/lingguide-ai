"""知识库管理 API"""
from fastapi import APIRouter, UploadFile, File, Depends, HTTPException
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession
from pydantic import BaseModel, Field
from typing import List
import os
import re
import unicodedata

from app.core.database import get_db
from app.core.rag import rag_service, SQLiteLexicalRetriever
from app.models import Document, Chunk, FAQ, IndexManifest, IndexJob
from app.api.dependencies import require_admin
from app.config import resolve_runtime_path, settings
from app.services.knowledge_service import IndexNotReadyError, KnowledgeService

router = APIRouter(prefix="/api/knowledge", tags=["知识库"])


# ===== 文档管理 =====

class DocumentOut(BaseModel):
    id: str
    filename: str
    file_type: str
    status: str
    chunk_count: int
    uploaded_at: str
    error_message: str = ""

    model_config = dict(from_attributes=True)


@router.get("/documents", response_model=List[DocumentOut])
async def list_documents(
    db: AsyncSession = Depends(get_db),
    _admin=Depends(require_admin),
):
    """获取 SQLite 中的真实文档状态。"""
    rows = (await db.execute(select(Document).order_by(Document.uploaded_at.desc()))).scalars().all()
    return [
        DocumentOut(
            id=row.id,
            filename=row.filename,
            file_type=row.file_type or "unknown",
            status=row.status or "uploaded",
            chunk_count=row.chunk_count or 0,
            uploaded_at=row.uploaded_at.isoformat() if row.uploaded_at else "",
            error_message=row.error_message or "",
        )
        for row in rows
    ]


@router.post("/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    _admin=Depends(require_admin),
):
    """上传知识文档并同步写入 SQLite、FTS5 和 Chroma。"""
    content = await file.read()
    try:
        service = KnowledgeService(resolve_runtime_path(settings.upload_dir))
        document = await service.ingest(db, file.filename or "", content)
    except IndexNotReadyError as exc:
        raise HTTPException(503, str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc
    except Exception as exc:
        # 不把内部路径、凭据或堆栈暴露给管理端；详细原因已写入 Document.error_message。
        raise HTTPException(500, "文档索引失败，请查看文档状态中的错误信息") from exc
    return {
        "id": document.id,
        "filename": document.filename,
        "file_type": document.file_type,
        "status": document.status,
        "chunk_count": document.chunk_count,
    }


@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    db: AsyncSession = Depends(get_db),
    _admin=Depends(require_admin),
):
    """按稳定文档 ID 删除知识文档及索引。"""
    deleted = await KnowledgeService(resolve_runtime_path(settings.upload_dir)).delete(db, document_id)
    if not deleted:
        raise HTTPException(404, "文档不存在")
    return {"deleted": document_id}


# ===== FAQ 管理 =====

import json as _json

FAQ_FILE = resolve_runtime_path(settings.faq_path)

# 默认 FAQ（首次启动时初始化）
DEFAULT_FAQS = [
    {
        "id": 1,
        "question": "灵山大佛有多高？",
        "answer": "灵山大佛高达88米，加上基座总高度达101.5米，是目前世界上最高的青铜立佛像之一。大佛由1560块青铜板拼焊而成，总用铜量达700多吨，于1997年落成开光。",
        "match_text": "大佛 多高 高度",
        "tags": ["景点"],
    },
    {
        "id": 2,
        "question": "梵宫有什么特色？",
        "answer": "灵山梵宫被誉为东方卢浮宫，是灵山胜境的核心建筑。它以佛教艺术为主题，穹顶壁画《天象图》精美绝伦，融合了东阳木雕、琉璃壁画、铜雕等多种传统工艺与现代建筑美学。",
        "match_text": "梵宫 特色 特点",
        "tags": ["景点", "建筑"],
    },
    {
        "id": 3,
        "question": "九龙灌浴表演时间是？",
        "answer": "九龙灌浴表演每天有四场：上午10:00、11:30，下午14:00、15:30。每场表演约15分钟，是灵山胜境最受欢迎的景观之一。建议提前10分钟到达以获得最佳观赏位置。",
        "match_text": "九龙 表演 时间",
        "tags": ["景点", "时间"],
    },
    {
        "id": 4,
        "question": "景区开放时间和门票价格？",
        "answer": "灵山胜境景区全年开放，时间为每日08:00-17:30。旺季门票约210元/人，淡季约160元/人。60岁以上老人和学生凭证件可享半价优惠，1.2米以下儿童免票。建议提前在官方渠道购票。",
        "match_text": "开放时间 门票 价格 多少钱",
        "tags": ["票务", "时间"],
    },
    {
        "id": 5,
        "question": "灵山有什么好吃的推荐？",
        "answer": "灵山梵宫内设有素斋餐厅，招牌菜品包括罗汉斋和灵山素面，人均消费约50-80元。此外还有素食小吃摊位，提供素包子、素糕点等特色小吃，口味清淡雅致。",
        "match_text": "好吃 美食 素斋 餐饮 推荐",
        "tags": ["餐饮"],
    },
    {
        "id": 6,
        "question": "灵山大佛是什么时候建造的？",
        "answer": "灵山大佛于1994年奠基，1997年11月15日落成开光。由中国佛教协会会长赵朴初提出五方五佛理念，灵山大佛作为东方大佛填补了空缺。大佛由1560块青铜板拼焊而成，总用铜量达700多吨。",
        "match_text": "建造 建造于 什么时候 何时",
        "tags": ["景点", "历史"],
    },
    {
        "id": 7,
        "question": "五印坛城是什么？",
        "answer": "五印坛城是灵山胜境的重要景点，是一座藏传佛教风格的建筑。坛城建筑精美，展示了藏传佛教的曼荼罗艺术，内部供奉有佛像和唐卡，游客可以近距离感受藏传佛教文化的庄严与神秘。",
        "match_text": "五印坛城 坛城 藏传",
        "tags": ["景点", "建筑"],
    },
    {
        "id": 8,
        "question": "曼飞龙塔有什么特点？",
        "answer": "曼飞龙塔是灵山胜境的标志性建筑之一，具有浓郁的南传佛教风格。塔身洁白，造型独特，是拍照打卡的热门地点。塔群错落有致，展现了东南亚佛教建筑的艺术魅力。",
        "match_text": "曼飞龙塔 白塔",
        "tags": ["景点", "建筑"],
    },
    {
        "id": 9,
        "question": "天下第一掌是什么？",
        "answer": "天下第一掌是灵山胜境的一处著名景观，为一只巨大的铜铸手掌。游客常在此摸掌祈福，寓意五福临门。铜掌按灵山大佛右手原样放大铸造，高11.7米，宽5.5米。",
        "match_text": "天下第一掌 铜掌 手掌",
        "tags": ["景点"],
    },
    {
        "id": 10,
        "question": "祥符禅寺有什么历史？",
        "answer": "祥符禅寺始建于唐代，已有千年历史，是灵山胜境的核心寺庙。寺内建筑古朴庄严，是佛教禅宗的重要道场，也是游客礼佛参拜的主要场所。寺前有古银杏广场，秋日金黄满地。",
        "match_text": "祥符禅寺 寺庙 历史",
        "tags": ["景点", "历史"],
    },
    {
        "id": 11,
        "question": "怎么去灵山胜境？交通方便吗？",
        "answer": "从无锡市区可乘坐88路、89路公交直达灵山胜境，车程约40分钟。也可乘坐地铁2号线至梅园站后换乘公交。自驾游客导航至灵山胜境停车场即可，停车场容量充足。",
        "match_text": "交通 公交 怎么去 自驾",
        "tags": ["交通"],
    },
    {
        "id": 12,
        "question": "游览灵山胜境需要多长时间？",
        "answer": "灵山胜境景区面积较大，建议预留4-5小时游览。主要景点包括灵山大佛、梵宫、九龙灌浴、五印坛城、曼飞龙塔、祥符禅寺等。如赶时间，核心景点2-3小时也可游览完毕。",
        "match_text": "游览 多长时间 多久 几小时",
        "tags": ["攻略"],
    },
    {
        "id": 13,
        "question": "灵山胜境适合带老人小孩去吗？",
        "answer": "灵山胜境非常适合全家出游。景区设有无障碍通道，大佛基座有电梯可直达，方便老人游览。梵宫文化体验馆有亲子手作活动，九龙灌浴的壮观表演也深受小朋友喜爱。",
        "match_text": "老人 小孩 儿童 亲子 轮椅 无障碍 奶奶 爷爷 长辈",
        "tags": ["攻略"],
    },
    {
        "id": 14,
        "question": "灵山的历史文化",
        "answer": "灵山胜境位于无锡太湖之滨，佛教文化源远流长。唐代祥符禅寺的建立奠定了其佛教地位，1990年代灵山大佛的建造使之成为世界级佛教圣地。景区融合了汉传、藏传、南传三大佛教流派文化。",
        "match_text": "历史 文化 佛教",
        "tags": ["文化", "历史"],
    },
    {
        "id": 15,
        "question": "灵山有哪些值得买的纪念品？",
        "answer": "灵山胜境有多处文创商店，推荐购买灵山福牌、梵宫书签、佛像小摆件、素斋调料包等特色纪念品。景区还提供祈福带和祈福牌，游客可书写心愿后挂在指定区域。",
        "match_text": "纪念品 购物 买 文创",
        "tags": ["购物"],
    },
]

def _normalize_faq_term(value: str) -> str:
    """统一 FAQ 规则和用户问题的文本格式。"""
    text = unicodedata.normalize("NFKC", str(value or "")).casefold().strip()
    return re.sub(r"[\s\W_]+", "", text, flags=re.UNICODE)


def _unique_terms(values: list[str]) -> list[str]:
    """清理 FAQ 词项并按首次出现顺序去重。"""
    result = []
    seen = set()
    for value in values or []:
        value = str(value).strip()
        key = _normalize_faq_term(value)
        if key and key not in seen:
            seen.add(key)
            result.append(value)
    return result


def _normalize_faq_entry(entry: dict) -> dict:
    """补齐 FAQ 字段并清理词项，避免历史数据无法参与联合匹配。"""
    item = dict(entry)
    item["question"] = str(item.get("question", "")).strip()
    item["answer"] = str(item.get("answer", "")).strip()
    item["match_text"] = str(item.get("match_text") or "").strip()
    item["tags"] = _unique_terms(item.get("tags", []))
    item["entities"] = _unique_terms(item.get("entities", []))
    item["intent"] = str(item.get("intent") or "").strip()
    item["intent_keywords"] = _unique_terms(item.get("intent_keywords", []))
    item["exact_questions"] = _unique_terms(item.get("exact_questions", []))
    return item


def _validate_faq_entry(entry: dict) -> None:
    """验证 FAQ 运行数据必须具备联合匹配所需字段。"""
    required = ("question", "answer", "entities", "intent", "intent_keywords")
    if any(not entry.get(field) for field in required):
        raise ValueError(f"FAQ 字段不完整: {entry.get('id', 'unknown')}")


def _load_faqs() -> List[dict]:
    data = None
    if os.path.exists(FAQ_FILE):
        try:
            with open(FAQ_FILE, "r", encoding="utf-8") as f:
                data = _json.load(f)
        except Exception:
            pass
    if not isinstance(data, list):
        data = []
        for item in DEFAULT_FAQS:
            legacy = dict(item)
            legacy["entities"] = [legacy.get("question", "").rstrip("？?!！ ")[:8]]
            legacy["intent"] = "general_intro"
            legacy["intent_keywords"] = str(legacy.get("match_text") or "").split()
            data.append(legacy)
    normalized = [_normalize_faq_entry(item) for item in data if isinstance(item, dict)]
    try:
        for item in normalized:
            _validate_faq_entry(item)
    except ValueError as exc:
        raise RuntimeError(f"FAQ 数据校验失败: {exc}") from exc
    if normalized != data:
        _save_faqs(normalized)
    return normalized


def _save_faqs(data: List[dict]):
    os.makedirs(os.path.dirname(FAQ_FILE), exist_ok=True)
    with open(FAQ_FILE, "w", encoding="utf-8") as f:
        _json.dump(data, f, ensure_ascii=False, indent=2)


FAQ_LIST: List[dict] = _load_faqs()


class FAQCreate(BaseModel):
    question: str = Field(min_length=1)
    answer: str = Field(min_length=1)
    match_text: str = ""
    tags: List[str] = Field(default_factory=list)
    entities: List[str] = Field(min_length=1)
    intent: str = Field(min_length=1)
    intent_keywords: List[str] = Field(min_length=1)
    exact_questions: List[str] = Field(default_factory=list)


def _prepare_faq(faq: FAQCreate) -> dict:
    """规范化 API 输入，确保空白词项不会绕过字段校验。"""
    entry = _normalize_faq_entry(faq.model_dump())
    try:
        _validate_faq_entry(entry)
    except ValueError as exc:
        raise HTTPException(422, "问题、答案、实体、意图和意图关键词均不能为空") from exc
    return entry


def _find_faq_conflicts(candidate: dict, existing_faqs: list[dict]) -> list[dict]:
    """检测重复问题及不同意图间无法区分的实体/关键词规则。"""
    conflicts = []
    candidate_questions = {_normalize_faq_term(candidate.get("question", ""))}
    candidate_questions.update(_normalize_faq_term(q) for q in candidate.get("exact_questions", []))
    candidate_entities = [_normalize_faq_term(x) for x in candidate.get("entities", [])]
    candidate_intents = [_normalize_faq_term(x) for x in candidate.get("intent_keywords", [])]
    for existing in existing_faqs:
        if existing.get("id") == candidate.get("id"):
            continue
        existing_questions = {_normalize_faq_term(existing.get("question", ""))}
        existing_questions.update(_normalize_faq_term(q) for q in existing.get("exact_questions", []))
        shared_questions = sorted(q for q in candidate_questions & existing_questions if q)
        shared_entities = sorted({c for c in candidate_entities for e in existing.get("entities", [])
                                  if c and (c in _normalize_faq_term(e) or _normalize_faq_term(e) in c)})
        shared_keywords = sorted({c for c in candidate_intents for e in existing.get("intent_keywords", [])
                                  if c and (c in _normalize_faq_term(e) or _normalize_faq_term(e) in c)})
        different_intent = _normalize_faq_term(candidate.get("intent")) != _normalize_faq_term(existing.get("intent"))
        if shared_questions or (different_intent and shared_entities and shared_keywords):
            conflicts.append({
                "existing_id": existing.get("id"),
                "shared_questions": shared_questions,
                "shared_entities": shared_entities,
                "shared_intent_keywords": shared_keywords,
                "existing_intent": existing.get("intent", ""),
                "candidate_intent": candidate.get("intent", ""),
            })
    return conflicts


@router.get("/faqs")
async def list_faqs(_admin=Depends(require_admin)):
    return FAQ_LIST


@router.post("/faqs")
async def create_faq(faq: FAQCreate, _admin=Depends(require_admin)):
    entry = _prepare_faq(faq)
    entry["id"] = max([f.get("id", 0) for f in FAQ_LIST], default=0) + 1
    conflicts = _find_faq_conflicts(entry, FAQ_LIST)
    if conflicts:
        raise HTTPException(409, {"code": "FAQ_CONFLICT", "message": "FAQ 规则与已有记录冲突", "conflicts": conflicts})
    FAQ_LIST.append(entry)
    _save_faqs(FAQ_LIST)
    return entry


@router.put("/faqs/{faq_id}")
async def update_faq(faq_id: int, faq: FAQCreate, _admin=Depends(require_admin)):
    entry = _prepare_faq(faq)
    index = next((i for i, item in enumerate(FAQ_LIST) if item.get("id") == faq_id), None)
    if index is None:
        raise HTTPException(404, "FAQ 不存在")
    entry["id"] = faq_id
    conflicts = _find_faq_conflicts(entry, FAQ_LIST)
    if conflicts:
        raise HTTPException(409, {"code": "FAQ_CONFLICT", "message": "FAQ 规则与已有记录冲突", "conflicts": conflicts})
    FAQ_LIST[index] = entry
    _save_faqs(FAQ_LIST)
    return entry


@router.delete("/faqs/{faq_id}")
async def delete_faq(faq_id: int, _admin=Depends(require_admin)):
    global FAQ_LIST
    if not any(f.get("id") == faq_id for f in FAQ_LIST):
        raise HTTPException(404, "FAQ 不存在")
    FAQ_LIST = [f for f in FAQ_LIST if f["id"] != faq_id]
    _save_faqs(FAQ_LIST)
    return {"deleted": faq_id}


# ===== 知识库状态 =====

@router.get("/stats")
async def knowledge_stats(
    db: AsyncSession = Depends(get_db),
    _admin=Depends(require_admin),
):
    """返回当前 active manifest 对应的知识库统计。"""
    from app.core.database import _sqlite_path
    from app.core.index_runtime import get_active_index
    import sqlite3

    active = get_active_index(_sqlite_path)
    canonical_count = await db.scalar(
        select(func.count(Chunk.id)).where(Chunk.status == "ready")
    ) or 0
    lexical = SQLiteLexicalRetriever(_sqlite_path, namespace=active["fts_namespace"])
    fts_count = 0
    if lexical.available:
        try:
            with sqlite3.connect(lexical.db_path) as conn:
                fts_count = int(conn.execute(
                    f"SELECT COUNT(*) FROM {active['fts_namespace']}"
                ).fetchone()[0])
        except sqlite3.Error:
            fts_count = 0
    vector_count = 0
    try:
        if rag_service.vector:
            collection = rag_service.vector.client.get_collection(active["vector_collection"])
            vector_count = int(collection.count())
    except Exception:
        vector_count = 0
    return {
        "chunk_count": int(canonical_count),
        "canonical_chunk_count": int(canonical_count),
        "fts_count": fts_count,
        "vector_count": vector_count,
        "orphan_vector_count": max(0, vector_count - int(canonical_count)),
        "faq_count": len(FAQ_LIST),
        "index_version": active["version"],
        "manifest_id": active.get("manifest_id"),
        "fts_namespace": active["fts_namespace"],
        "vector_collection": active["vector_collection"],
    }


# ===== 工具函数 =====

def _parse_document(file_path: str, ext: str) -> str:
    """解析文档为纯文本"""
    if ext == "txt" or ext == "md":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    elif ext == "docx":
        import docx
        doc = docx.Document(file_path)
        parts = []

        # 1. 提取段落文本
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        # 2. 提取表格数据（每行格式化为一条结构化记录）
        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                if not any(cells):
                    continue
                # 跳过标题行重复
                if cells[0] == headers[0]:
                    continue

                # 景区结构化表格（11列）：景区名称 | 景点ID | 景点名称 | 位置 | 参数 | 功能 | 文化 | 介绍 | 亮点 | 开放 | 备注
                if len(headers) >= 8 and '景点名称' in headers:
                    idx_name = headers.index('景点名称') if '景点名称' in headers else 0
                    idx_loc = headers.index('具体位置') if '具体位置' in headers else 3
                    idx_detail = headers.index('详细介绍') if '详细介绍' in headers else 7
                    idx_highlight = headers.index('游玩亮点') if '游玩亮点' in headers else 8
                    idx_culture = headers.index('文化内涵') if '文化内涵' in headers else 6
                    idx_open = headers.index('演艺/开放信息') if '演艺/开放信息' in headers else 9

                    record = f"【{cells[idx_name]}】\n"
                    if idx_loc < len(cells): record += f"位置：{cells[idx_loc]}\n"
                    if idx_detail < len(cells): record += f"介绍：{cells[idx_detail]}\n"
                    if idx_highlight < len(cells): record += f"亮点：{cells[idx_highlight]}\n"
                    if idx_culture < len(cells): record += f"文化内涵：{cells[idx_culture]}\n"
                    if idx_open < len(cells): record += f"开放信息：{cells[idx_open]}\n"
                    parts.append(record)
                else:
                    # 通用表格：用 tab 拼接
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts)

    elif ext == "pdf":
        # 简化版；完整版用 pypdf 或 pdfplumber
        raise HTTPException(400, "PDF 解析请安装 pypdf 依赖")

    return ""


def _split_text(text: str, chunk_size: int = 512, overlap: int = 64) -> List[str]:
    """简单滑动窗口切片（中文友好版）"""
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks
